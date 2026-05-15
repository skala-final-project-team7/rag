"""첨부 3유형 청킹 검증 (feature4-A: docx / xlsx) — chunking-strategy.md §5.

split_attachment는 첨부 파일을 attachment_type별 전략으로 1차 분할하고,
chunk_attachment는 크기 규칙·메타데이터 부착까지 거쳐 Chunk 목록을 산출한다.
실제 픽스처는 samples/attachments/ 의 docx 2건·xlsx 2건을 사용한다.
"""

from datetime import datetime
from pathlib import Path

import openpyxl
import pytest
from docx import Document as DocxDocument

from app.adapters.json_fixture import JsonFixtureSourceAdapter
from app.ingestion.chunker.attachment import (
    build_attachment_metadata,
    chunk_attachment,
    infer_attachment_type,
    split_attachment,
)
from app.ingestion.chunker.base import ChunkDraft
from app.schemas.chunk import Chunk, make_chunk_id
from app.schemas.enums import AttachmentType, ExtractedFormat, SourceType
from app.schemas.page_object import Attachment, PageObject

SAMPLES_DIR = Path(__file__).resolve().parents[3] / "samples"
ATTACHMENTS_DIR = SAMPLES_DIR / "attachments"

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

_PARENT_PAGE = PageObject(
    page_id="CONF-PAGE-1",
    space_key="CLOUD",
    title="EKS 장애 대응 가이드",
    body_html="<h2>본문</h2>",
    version_number=3,
    last_modified="2026-04-22T08:15:00+09:00",
    allowed_groups=["space:CLOUD"],
    allowed_users=["user:taesung"],
    webui_link="/display/CLOUD/eks",
    labels=["eks", "장애대응"],
    ancestors=["Cloud 운영 문서", "EKS 운영"],
)


def _attachment(filename: str, mime_type: str) -> Attachment:
    """samples/attachments/<filename>을 가리키는 Attachment 픽스처."""
    return Attachment(
        attachment_id="CONF-PAGE-1-att-0",
        filename=filename,
        mime_type=mime_type,
        extracted_text="",
        extracted_format=ExtractedFormat.RAW_TEXT,
        download_url=str(ATTACHMENTS_DIR / filename),
        parent_page_id="CONF-PAGE-1",
        last_modified=datetime.fromisoformat("2026-04-20T10:00:00+09:00"),
    )


_DOCX_MANUAL = _attachment("EKS_운영_상세_매뉴얼_v2.3.docx", _DOCX_MIME)
_DOCX_ONBOARD = _attachment("신규입사자_온보딩_체크리스트_2026.docx", _DOCX_MIME)
_XLSX_METRICS = _attachment("모니터링_메트릭_정의서_v1.4.xlsx", _XLSX_MIME)
_XLSX_USAGE = _attachment("EKS_노드_월간_사용량_통계_2026Q1.xlsx", _XLSX_MIME)


# --- infer_attachment_type ---


def test_infer_attachment_type_from_mime() -> None:
    assert infer_attachment_type(_DOCX_MANUAL) is AttachmentType.DOCX
    assert infer_attachment_type(_XLSX_METRICS) is AttachmentType.XLSX


def test_infer_attachment_type_falls_back_to_extension() -> None:
    # mime이 generic이어도 확장자로 판별한다
    generic = _attachment("EKS_운영_상세_매뉴얼_v2.3.docx", "application/octet-stream")
    assert infer_attachment_type(generic) is AttachmentType.DOCX


def test_infer_attachment_type_rejects_unknown() -> None:
    unknown = _attachment("memo.hwp", "application/x-hwp")
    with pytest.raises(ValueError, match="첨부 유형"):
        infer_attachment_type(unknown)


# --- docx 1차 분할 (split_attachment) ---


def test_docx_splits_by_heading_hierarchy() -> None:
    drafts = split_attachment(_DOCX_MANUAL, AttachmentType.DOCX)
    assert all(isinstance(d, ChunkDraft) for d in drafts)
    # Heading 1/2/3 각각이 섹션 경계 — 픽스처는 헤딩 44개
    assert len(drafts) == 44
    headers = {d.section_header for d in drafts}
    assert {"0. 개정 이력", "1.1 문서 목적", "4.2 노드 조인 실패"} <= headers
    # 첨부 섹션은 원자성 없음 (2차 재분할·하한선 병합 대상)
    assert all(d.is_atomic is False for d in drafts)


def test_docx_prepends_preamble_to_first_section() -> None:
    drafts = split_attachment(_DOCX_MANUAL, AttachmentType.DOCX)
    first = drafts[0]
    assert first.section_header == "0. 개정 이력"
    # 첫 헤딩 이전 표지 문단(preamble)이 첫 섹션 도입부에 부착된다
    assert first.text.startswith("EKS 운영 상세 매뉴얼")
    assert "주 담당: 최태성, 신유진" in first.text


def test_docx_converts_table_to_markdown() -> None:
    drafts = split_attachment(_DOCX_MANUAL, AttachmentType.DOCX)
    first = drafts[0]
    # '0. 개정 이력' 섹션에 포함된 표가 markdown으로 변환된다
    assert "| 버전 | 일자 | 주요 변경 내용 | 작성자 |" in first.text
    assert "| --- | --- | --- | --- |" in first.text


def test_docx_headingless_falls_back_to_single_draft(tmp_path: Path) -> None:
    # 헤딩이 없는 docx는 단일 draft로 폴백하고 section_header는 파일명을 쓴다
    document = DocxDocument()
    document.add_paragraph("헤딩 없는 첫 문단입니다.")
    document.add_paragraph("헤딩 없는 둘째 문단입니다.")
    path = tmp_path / "plain.docx"
    document.save(path)
    plain = _attachment("plain.docx", _DOCX_MIME)
    plain = plain.model_copy(update={"download_url": str(path)})

    drafts = split_attachment(plain, AttachmentType.DOCX)
    assert len(drafts) == 1
    assert drafts[0].section_header == "plain.docx"
    assert "헤딩 없는 첫 문단입니다." in drafts[0].text
    assert "헤딩 없는 둘째 문단입니다." in drafts[0].text


# --- xlsx 1차 분할 (split_attachment) ---


def test_xlsx_splits_by_sheet_and_serializes_rows() -> None:
    drafts = split_attachment(_XLSX_METRICS, AttachmentType.XLSX)
    headers = {d.section_header for d in drafts}
    # 시트 단위 분할 — section_header는 '[시트명] 행 N~M'
    assert "[클러스터 메트릭] 행 1~10" in headers
    assert "[개정 이력] 행 1~4" in headers
    # 각 행이 '[<시트명>] <컬럼>: <값> | ...' 형식으로 직렬화된다 (컬럼명 매 행 동봉)
    cluster = next(d for d in drafts if d.section_header == "[클러스터 메트릭] 행 1~10")
    assert (
        "[클러스터 메트릭] 메트릭 ID: CL-001 | 메트릭 이름: kubernetes.node.cpu.usage.pct"
        in cluster.text
    )


def test_xlsx_omits_empty_cells(tmp_path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "표"
    sheet.append(["이름", "메모", "비고"])
    sheet.append(["alpha", None, "확인"])  # 가운데 셀 비어 있음
    path = tmp_path / "sparse.xlsx"
    workbook.save(path)
    sparse = _attachment("sparse.xlsx", _XLSX_MIME)
    sparse = sparse.model_copy(update={"download_url": str(path)})

    drafts = split_attachment(sparse, AttachmentType.XLSX)
    text = drafts[0].text
    # 빈 셀은 직렬화에서 생략된다
    assert "이름: alpha" in text
    assert "비고: 확인" in text
    assert "메모:" not in text


def test_xlsx_groups_rows_by_50(tmp_path: Path) -> None:
    # 직렬화 토큰이 작은 시트 → 50행 그룹 경계가 그대로 유지된다
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "수치"
    sheet.append(["idx", "value"])
    for i in range(1, 121):  # 데이터 120행
        sheet.append([i, i * 2])
    path = tmp_path / "rows.xlsx"
    workbook.save(path)
    rows_attach = _attachment("rows.xlsx", _XLSX_MIME)
    rows_attach = rows_attach.model_copy(update={"download_url": str(path)})

    drafts = split_attachment(rows_attach, AttachmentType.XLSX)
    headers = [d.section_header for d in drafts]
    assert headers == ["[수치] 행 1~50", "[수치] 행 51~100", "[수치] 행 101~120"]


def test_xlsx_oversized_group_shrinks_to_25(tmp_path: Path) -> None:
    # 50행 그룹 직렬화가 800토큰을 초과하면 25행으로 축소 재분할된다
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "긴행"
    sheet.append(["코드", "설명"])
    long_text = "토큰 수를 늘리기 위한 설명 문장"  # 행당 약 25토큰 → 50행 그룹이 800토큰 초과
    for i in range(1, 61):  # 데이터 60행
        sheet.append([f"C-{i}", long_text])
    path = tmp_path / "long.xlsx"
    workbook.save(path)
    long_attach = _attachment("long.xlsx", _XLSX_MIME)
    long_attach = long_attach.model_copy(update={"download_url": str(path)})

    drafts = split_attachment(long_attach, AttachmentType.XLSX)
    headers = [d.section_header for d in drafts]
    # 50행 그룹이 25행으로 축소: 1~25 / 26~50 / 51~60
    assert headers == ["[긴행] 행 1~25", "[긴행] 행 26~50", "[긴행] 행 51~60"]


def test_xlsx_synthesizes_header_when_missing(tmp_path: Path) -> None:
    # 첫 행이 데이터(수치)면 헤더 누락으로 보고 col_1, col_2... 를 부여한다 (ATTACH_NO_HEADER)
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "헤더없음"
    sheet.append([10, 20, 30])
    sheet.append([40, 50, 60])
    path = tmp_path / "noheader.xlsx"
    workbook.save(path)
    noheader = _attachment("noheader.xlsx", _XLSX_MIME)
    noheader = noheader.model_copy(update={"download_url": str(path)})

    drafts = split_attachment(noheader, AttachmentType.XLSX)
    text = drafts[0].text
    assert "col_1: 10" in text and "col_2: 20" in text and "col_3: 30" in text
    # 첫 행도 데이터로 포함된다
    assert "col_1: 40" in text


# --- 첨부 메타데이터 (build_attachment_metadata) ---


def test_attachment_metadata_inherits_from_parent_page() -> None:
    draft = ChunkDraft(text="[개정 이력] 버전: v1.0", section_header="[개정 이력] 행 1~4")
    meta = build_attachment_metadata(
        _PARENT_PAGE, _XLSX_METRICS, draft, chunk_index=0, attachment_type=AttachmentType.XLSX
    )
    # 첨부 전용 필드
    assert meta.source_type is SourceType.ATTACHMENT
    assert meta.attachment_id == "CONF-PAGE-1-att-0"
    assert meta.attachment_filename == "모니터링_메트릭_정의서_v1.4.xlsx"
    assert meta.attachment_mime == _XLSX_MIME
    assert meta.extracted_format is ExtractedFormat.SHEET_SERIALIZED
    # doc_type 필드는 첨부의 attachment_type 값을 담는다
    assert meta.doc_type == "xlsx"
    # ACL·페이지 메타는 부모 페이지에서 상속
    assert meta.page_id == "CONF-PAGE-1"
    assert meta.allowed_groups == ["space:CLOUD"]
    assert meta.allowed_users == ["user:taesung"]
    assert meta.labels == ["eks", "장애대응"]
    assert meta.space_key == "CLOUD"
    assert meta.webui_link == "/display/CLOUD/eks"
    assert meta.token_count > 0


def test_attachment_chunk_id_uses_attachment_id() -> None:
    draft = ChunkDraft(text="본문", section_header="섹션")
    meta = build_attachment_metadata(
        _PARENT_PAGE, _DOCX_MANUAL, draft, chunk_index=0, attachment_type=AttachmentType.DOCX
    )
    # chunk_id는 attachment_id를 포함한 결정론적 SHA1
    assert meta.chunk_id == make_chunk_id("CONF-PAGE-1", 0, "CONF-PAGE-1-att-0")
    # 본문 청크(attachment_id 없음)와 다른 id
    assert meta.chunk_id != make_chunk_id("CONF-PAGE-1", 0)


def test_attachment_metadata_extracted_format_for_docx() -> None:
    draft = ChunkDraft(text="본문", section_header="섹션")
    meta = build_attachment_metadata(
        _PARENT_PAGE, _DOCX_MANUAL, draft, chunk_index=0, attachment_type=AttachmentType.DOCX
    )
    assert meta.extracted_format is ExtractedFormat.RAW_TEXT
    assert meta.doc_type == "docx"


# --- chunk_attachment 엔트리 ---


def test_chunk_attachment_docx_returns_indexed_chunks() -> None:
    chunks = chunk_attachment(_DOCX_ONBOARD, _PARENT_PAGE, attachment_type=AttachmentType.DOCX)
    assert len(chunks) >= 1
    assert all(isinstance(c, Chunk) for c in chunks)
    assert [c.metadata.chunk_index for c in chunks] == list(range(len(chunks)))
    assert all(c.metadata.source_type is SourceType.ATTACHMENT for c in chunks)
    assert all(c.metadata.extracted_format is ExtractedFormat.RAW_TEXT for c in chunks)
    assert all(c.text.strip() for c in chunks)


def test_chunk_attachment_xlsx_returns_indexed_chunks() -> None:
    chunks = chunk_attachment(_XLSX_USAGE, _PARENT_PAGE, attachment_type=AttachmentType.XLSX)
    assert len(chunks) >= 1
    assert [c.metadata.chunk_index for c in chunks] == list(range(len(chunks)))
    assert all(c.metadata.source_type is SourceType.ATTACHMENT for c in chunks)
    assert all(c.metadata.extracted_format is ExtractedFormat.SHEET_SERIALIZED for c in chunks)
    # 92행 시트가 50→25행 축소 분할되어 여러 청크로 나뉜다
    assert len(chunks) > 4


def test_chunk_attachment_infers_type_when_omitted() -> None:
    # attachment_type 미지정 → mime 기반 추정 (PoC)
    chunks = chunk_attachment(_XLSX_METRICS, _PARENT_PAGE)
    assert len(chunks) >= 1
    assert all(c.metadata.doc_type == "xlsx" for c in chunks)


def test_chunk_attachment_rejects_unsupported_type() -> None:
    # feature4-A는 docx/xlsx만 지원 — PDF/CSV는 feature4-B
    pdf = _attachment("report.pdf", "application/pdf")
    with pytest.raises(ValueError, match="feature4-B"):
        chunk_attachment(pdf, _PARENT_PAGE, attachment_type=AttachmentType.PDF)


# --- samples/attachments 통합 청킹 ---


def test_samples_attachments_chunk_without_error() -> None:
    adapter = JsonFixtureSourceAdapter(samples_dir=SAMPLES_DIR)
    pairs = [
        (page, attachment) for page in adapter.fetch_pages() for attachment in page.attachments
    ]
    # confluence 샘플의 첨부 4건 (docx 2 + xlsx 2)
    assert len(pairs) == 4

    format_counts = {ExtractedFormat.RAW_TEXT: 0, ExtractedFormat.SHEET_SERIALIZED: 0}
    for page, attachment in pairs:
        chunks = chunk_attachment(attachment, page)
        assert len(chunks) >= 1, f"attachment {attachment.filename} produced no chunks"
        for index, chunk in enumerate(chunks):
            meta = chunk.metadata
            assert chunk.text.strip()
            assert meta.chunk_index == index
            assert meta.section_header, "section_header must not be empty"
            assert meta.source_type is SourceType.ATTACHMENT
            assert meta.attachment_id == attachment.attachment_id
            assert meta.page_id == page.page_id
            assert meta.allowed_groups == page.allowed_groups
            assert meta.token_count > 0
        format_counts[chunks[0].metadata.extracted_format] += 1

    # docx 2건 → raw_text, xlsx 2건 → sheet_serialized
    assert format_counts[ExtractedFormat.RAW_TEXT] == 2
    assert format_counts[ExtractedFormat.SHEET_SERIALIZED] == 2
